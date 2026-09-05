#!/usr/bin/env python3
"""The pywebview Api class — bridges Python backend to JS frontend.

Each method is callable from JS as: await window.pywebview.api.method_name(args)
pywebview runs these on a worker thread; the UI stays responsive.
"""
import os
import sys
import json
import shutil
from pathlib import Path
from datetime import datetime
from typing import Any, Dict, List, Optional

# Fix Windows Unicode
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except (AttributeError, Exception):
    pass

from . import db, llm, tagger, safety
from .config import PROJECT_ROOT, FOLDERS, STATUSES
from .file_io import read_text_file, write_text_file
from .analysis_catalog import ANALYSIS_CATALOG
from .analyzers import (
    craft, voice_tense, characters, continuity, themes, editor,
    cadence, motifs, anchors, voice_dna, reader_perception
)
from .analysis_suite import run as suite_run
from .writer_intelligence import chapter_comparison
from .search import (
    search_by_tag, search_multi, find_tag_in_file,
    get_tag_coverage, get_all_values_for_tag
)
from .export import export_markdown, export_plain_text, export_docx, export_analysis_report
from . import settings as settings_module


class Api:
    """Exposes all Scribbler functionality to the pywebview JS frontend."""

    # ── STATUS & FILES ──────────────────────────────────────────────

    def get_status(self) -> dict:
        """Return app status — AI availability, version, file counts."""
        stats = db.get_stats()
        return {
            "ok": True,
            "version": "5.0",
            "llm": llm.llm_status(),
            "total_files": stats.get("total_files", 0),
            "total_words": stats.get("total_words", 0),
        }

    def list_files(self) -> dict:
        """Return all files with metadata."""
        all_files = db.get_all_files()
        # Also scan folders for unindexed files
        seen = set()
        out = []
        for f in all_files:
            p = f.get("path")
            seen.add(p)
            out.append({
                "path": p,
                "filename": f.get("filename", ""),
                "folder": f.get("folder", ""),
                "word_count": f.get("word_count", 0),
                "status": f.get("status", "seedling"),
                "last_analyzed": f.get("last_analyzed", ""),
                "characters": f.get("characters", []),
                "themes": f.get("themes", []),
                "era": f.get("era", ""),
                "voice": f.get("voice", ""),
                "emotional_register": f.get("emotional_register", ""),
            })
        for folder in ("raw-dumps", "triage", "chapters", "drafts", "final"):
            root = PROJECT_ROOT / folder
            if root.exists():
                for p in root.iterdir():
                    if p.is_file() and p.suffix.lower() in (".txt", ".md", ".text") and p.name.upper() != "README.MD":
                        resolved = str(p.resolve())
                        if resolved not in seen:
                            seen.add(resolved)
                            try:
                                text = read_text_file(p)
                                wc = len(text.split())
                            except Exception:
                                wc = 0
                            out.append({
                                "path": str(p),
                                "filename": p.name,
                                "folder": folder,
                                "word_count": wc,
                                "status": "unindexed",
                                "last_analyzed": "",
                                "characters": [],
                                "themes": [],
                                "era": "",
                                "voice": "",
                                "emotional_register": "",
                            })
        return {"files": sorted(out, key=lambda x: (x["folder"], x["filename"].lower()))}

    def get_tools(self) -> dict:
        """Return the 17-tool analysis catalog."""
        return {
            "tools": {k: {"title": v[0], "group": v[1], "purpose": v[2]} for k, v in _get_tools_dict().items()},
            "catalog": dict(ANALYSIS_CATALOG),
        }

    # ── IMPORT & FILES ──────────────────────────────────────────────

    def import_files(self, destination: str, file_paths: list) -> dict:
        """Copy files into a project folder. file_paths come from native file dialog."""
        if destination not in ("raw-dumps", "triage", "chapters", "drafts", "final"):
            return {"ok": False, "error": f"Invalid destination: {destination}"}
        dest_folder = PROJECT_ROOT / destination
        dest_folder.mkdir(parents=True, exist_ok=True)
        count = 0
        errors = []
        for fp in file_paths:
            try:
                src = Path(fp)
                if not src.exists():
                    errors.append(f"{fp}: file not found")
                    continue
                name = _safe_name(src.name)
                dest = _unique_path(dest_folder, name)
                shutil.copy2(str(src), str(dest))
                count += 1
            except Exception as e:
                errors.append(f"{fp}: {e}")
        return {"ok": True, "message": f"Imported {count} file(s) into {destination}", "errors": errors}

    def save_note(self, title: str, text: str) -> dict:
        """Save a quick note to raw-dumps."""
        if not text.strip():
            return {"ok": False, "error": "Note is empty"}
        name = _safe_name((title.strip() or f"note-{datetime.now():%Y%m%d-%H%M%S}") + ".txt")
        dest = _unique_path(PROJECT_ROOT / "raw-dumps", name)
        write_text_file(dest, text)
        return {"ok": True, "message": "Saved to Inbox"}

    def delete_file(self, path: str) -> dict:
        """Move a file to archive."""
        try:
            p = Path(path)
            if not p.exists():
                return {"ok": False, "error": "File not found"}
            archive = PROJECT_ROOT / "archive"
            archive.mkdir(exist_ok=True)
            dest = _unique_path(archive, p.name)
            p.rename(dest)
            # Delete from DB — try both the original path and resolved path
            conn = db.get_db()
            resolved = str(p.resolve())
            original = str(p)
            # Try multiple path variations the DB might store
            for path_var in [resolved, original, path]:
                conn.execute("DELETE FROM files WHERE path = ?", (path_var,))
                conn.execute("DELETE FROM analysis_results WHERE file_path = ?", (path_var,))
            # Also try LIKE match for relative/absolute mismatches
            conn.execute("DELETE FROM files WHERE path LIKE ?", (f"%{p.name}%",))
            conn.execute("DELETE FROM analysis_results WHERE file_path LIKE ?", (f"%{p.name}%",))
            conn.execute("INSERT INTO activity_log (timestamp, action, file_path, details) VALUES (?, ?, ?, ?)",
                         (datetime.now().isoformat(), "delete", original, f"Moved to archive/{dest.name}"))
            conn.commit()
            conn.close()
            return {"ok": True, "message": f"Moved to archive/{dest.name}"}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    # ── TAGGING ─────────────────────────────────────────────────────

    def tag_preview(self, paths: list, use_ai: bool = False) -> dict:
        """Preview tags for selected files without applying."""
        if not paths:
            return {"ok": False, "error": "Select one or more files first"}
        previews = []
        errors = []
        for raw_path in paths:
            try:
                p = Path(raw_path)
                text = read_text_file(p)
                # Strip frontmatter
                if text.startswith("---"):
                    end = text.find("---", 3)
                    if end != -1:
                        text = text[end + 3:].strip()
                import re
                text = re.sub(r'<!-- SCRIBBLER SUMMARY[\s\S]*?-->', '', text).strip()
                previews.append({
                    "filename": p.name,
                    "path": str(p),
                    "word_count": len(text.split()),
                    "voice": tagger.detect_voice(text),
                    "era": tagger.detect_era(text),
                    "emotional_register": tagger.detect_emotional_register(text),
                    "sensory": tagger.detect_sensory(text),
                    "themes": tagger.detect_themes(text),
                    "characters": tagger.detect_characters(text),
                    "places": tagger.detect_places(text),
                })
            except Exception as e:
                errors.append(f"{raw_path}: {e}")
        return {"ok": True, "preview": previews, "errors": errors}

    def tag_files(self, paths: list, use_llm: bool = True) -> dict:
        """Apply tags to selected files. Pushes progress via evaluate_js."""
        if not paths:
            return {"ok": False, "error": "Select one or more files first"}
        tagged = []
        errors = []
        total = len(paths)
        for i, raw_path in enumerate(paths):
            # Push progress to JS
            self._push_progress(i + 1, total, f"Tagging {Path(raw_path).name}")
            try:
                meta = tagger.tag_file(raw_path, use_llm=use_llm)
                tagged.append(Path(raw_path).name)
            except Exception as e:
                errors.append(f"{Path(raw_path).name}: {e}")
        self._push_progress(total, total, "Done")
        return {"ok": True, "tagged": tagged, "errors": errors}

    # ── ANALYSIS ────────────────────────────────────────────────────

    def analyze(self, paths: list, tools: list) -> dict:
        """Run analysis tools on selected files. Pushes step progress."""
        if not paths:
            return {"ok": False, "error": "Select one or more manuscript files first"}
        if not tools:
            return {"ok": False, "error": "Choose at least one analysis tool"}

        all_files = self.list_files().get("files", [])
        results = []
        total_steps = len(paths) * len(tools)
        step = 0

        for raw_path in paths:
            try:
                p = Path(raw_path)
                text = read_text_file(p)
                # Strip frontmatter
                if text.startswith("---"):
                    end = text.find("---", 3)
                    if end != -1:
                        text = text[end + 3:].strip()
                import re
                text = re.sub(r'<!-- SCRIBBLER SUMMARY[\s\S]*?-->', '', text).strip()

                per_file = {}
                for tool_key in tools:
                    step += 1
                    self._push_progress(step, total_steps, f"Analysing {p.name} · {tool_key}")
                    try:
                        result = _run_tool(tool_key, text, all_files)
                        per_file[tool_key] = _js_safe(result)
                        try:
                            db.save_analysis(str(p.resolve()), tool_key, result)
                        except Exception:
                            pass
                    except Exception as e:
                        per_file[tool_key] = {"error": str(e)}
                results.append({"filename": p.name, "results": per_file})
            except Exception as e:
                results.append({"filename": raw_path, "results": {}, "error": str(e)})

        self._push_progress(total_steps, total_steps, "Done")
        return {"ok": True, "results": results, "message": f"Analysed {len(paths)} file(s) with {len(tools)} tool(s)"}

    def compare_chapters(self, paths: list) -> dict:
        """Run cross-chapter comparison (voice drift, pacing, tone)."""
        if not paths or len(paths) < 2:
            return {"ok": False, "error": "Select at least 2 chapters to compare"}
        chapters = []
        total = len(paths)
        for i, raw_path in enumerate(paths):
            self._push_progress(i + 1, total, f"Reading {Path(raw_path).name}")
            try:
                p = Path(raw_path)
                text = read_text_file(p)
                if text.startswith("---"):
                    end = text.find("---", 3)
                    if end != -1:
                        text = text[end + 3:].strip()
                chapters.append({"filename": p.name, "text": text})
            except Exception as e:
                return {"ok": False, "error": f"Could not read {raw_path}: {e}"}
        self._push_progress(total, total, "Comparing chapters")
        try:
            result = chapter_comparison(chapters)
            return {"ok": True, "result": _js_safe(result)}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    # ── SEARCH ──────────────────────────────────────────────────────

    def search_tags(self, tag_type: str, value: str) -> dict:
        """Search files by a single tag."""
        results = search_by_tag(tag_type, value)
        return {"ok": True, "results": _js_safe(results), "count": len(results)}

    def search_multi_tags(self, filters: dict) -> dict:
        """Search with multiple tag filters (AND logic)."""
        results = search_multi(filters)
        return {"ok": True, "results": _js_safe(results), "count": len(results)}

    def find_in_file(self, path: str, tag_type: str, value: str) -> dict:
        """Find where a tag value appears in a file (paragraph-level)."""
        occurrences = find_tag_in_file(path, tag_type, value)
        return {"ok": True, "occurrences": _js_safe(occurrences), "count": len(occurrences)}

    def tag_coverage(self, path: str) -> dict:
        """Get tag coverage report for a file."""
        coverage = get_tag_coverage(path)
        return {"ok": True, "coverage": _js_safe(coverage)}

    def get_tag_values(self, tag_type: str) -> dict:
        """Get all unique values for a tag type."""
        values = get_all_values_for_tag(tag_type)
        return {"ok": True, "values": _js_safe(values)}

    # ── EXPORT ──────────────────────────────────────────────────────

    def pick_open_files(self) -> dict:
        """Open a native file picker dialog. Returns selected file paths."""
        try:
            import webview
            window = webview.windows[0] if webview.windows else None
            if not window:
                return {"paths": []}
            result = window.create_file_dialog(
                webview.OPEN_DIALOG,
                allow_multiple=True,
                file_types=('Text Files (*.txt;*.md;*.text)', 'All Files (*.*)'),
            )
            if result:
                return {"paths": result if isinstance(result, list) else [result]}
            return {"paths": []}
        except Exception as e:
            return {"paths": [], "error": str(e)}

    def pick_save_path(self, default_name: str = "export.txt") -> dict:
        """Open a native save dialog. Returns chosen save path."""
        try:
            import webview
            window = webview.windows[0] if webview.windows else None
            if not window:
                return {"path": None}
            # Determine file extension filter
            ext = ".txt"
            if default_name.endswith(".docx"):
                ext = ".docx"
                file_types = ('Word Document (*.docx)',)
            elif default_name.endswith(".md"):
                ext = ".md"
                file_types = ('Markdown (*.md)',)
            elif default_name.endswith(".zip"):
                ext = ".zip"
                file_types = ('ZIP Archive (*.zip)',)
            else:
                file_types = ('Text Files (*.txt)',)

            result = window.create_file_dialog(
                webview.SAVE_DIALOG,
                save_filename=default_name,
                file_types=file_types,
            )
            if result:
                return {"path": result}
            return {"path": None}
        except Exception as e:
            return {"path": None, "error": str(e)}

    def export_file(self, path: str, kind: str, save_path: str) -> dict:
        """Export a file to docx/md/txt at a user-chosen location."""
        try:
            # save_path may come as a tuple from pywebview — extract the string
            if isinstance(save_path, (list, tuple)):
                save_path = save_path[0] if save_path else None
            if not save_path:
                return {"ok": False, "error": "No save location chosen"}

            # Generate the export to the user's chosen path directly
            from pathlib import Path as P
            dest = P(save_path)
            dest.parent.mkdir(parents=True, exist_ok=True)

            src = P(path)
            if not src.exists():
                return {"ok": False, "error": "Source file not found"}

            if kind == "docx":
                from .export import _sanitize_for_docx
                try:
                    from docx import Document
                    from docx.shared import Pt
                except ImportError:
                    return {"ok": False, "error": "python-docx not installed"}
                content = read_text_file(src)
                if content.startswith("---"):
                    end = content.find("---", 3)
                    if end != -1:
                        content = content[end + 3:].strip()
                import re as _re
                content = _re.sub(r'<!-- SCRIBBLER SUMMARY[\s\S]*?-->', '', content).strip()
                content = _sanitize_for_docx(content)
                doc = Document()
                style = doc.styles['Normal']
                style.font.name = 'Calibri'
                style.font.size = Pt(11)
                title = src.stem.replace('-', ' ').replace('_', ' ').title()
                title = _sanitize_for_docx(title)
                doc.add_heading(title, level=1)
                for para in _re.split(r'\n\s*\n', content):
                    para = para.strip()
                    if not para:
                        continue
                    if para.startswith('# '):
                        doc.add_heading(_sanitize_for_docx(para[2:]), level=1)
                    elif para.startswith('## '):
                        doc.add_heading(_sanitize_for_docx(para[3:]), level=2)
                    elif para.startswith('### '):
                        doc.add_heading(_sanitize_for_docx(para[4:]), level=3)
                    else:
                        doc.add_paragraph(_sanitize_for_docx(para))
                doc.save(str(dest))
            elif kind == "md":
                content = read_text_file(src)
                write_text_file(dest, content)
            elif kind == "txt":
                content = read_text_file(src)
                if content.startswith("---"):
                    end = content.find("---", 3)
                    if end != -1:
                        content = content[end + 3:].strip()
                import re as _re2
                content = _re2.sub(r'<!-- SCRIBBLER SUMMARY[\s\S]*?-->', '', content).strip()
                write_text_file(dest, content)
            else:
                return {"ok": False, "error": f"Unknown format: {kind}"}

            return {"ok": True, "path": str(dest)}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    def backup_project(self, save_path: str) -> dict:
        """Create a ZIP backup at a user-chosen location."""
        try:
            out = safety.export_project_zip()
            if save_path and save_path != out:
                shutil.move(out, save_path)
                out = save_path
            return {"ok": True, "path": out}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    # ── SETTINGS ────────────────────────────────────────────────────

    def get_ai_status(self) -> dict:
        return {"status": llm.llm_status(), "available": llm.llm_available()}

    def set_ai_provider(self, provider: str, api_key: str) -> dict:
        settings_module.set_setting("provider", provider)
        settings_module.set_setting("api_key", api_key)
        return {"ok": True, "message": f"Provider set to {provider}"}

    # ── INTERNAL HELPERS ────────────────────────────────────────────

    def _push_progress(self, step: int, total: int, message: str):
        """Push progress to the JS frontend via evaluate_js."""
        try:
            import webview
            for window in webview.windows:
                window.evaluate_js(
                    f"window.__scribblerProgress__({step}, {total}, '{message}')"
                )
        except Exception:
            pass  # In headless/test mode, no window exists


# ── MODULE-LEVEL HELPERS ────────────────────────────────────────────

def _get_tools_dict():
    """Return the TOOLS dict mapping tool keys to (title, group, purpose, fn)."""
    return {
        "craft": ("Craft & Rhythm", "Prose", "Sentence rhythm, balance and craft signals.", craft.analyze),
        "voice": ("Voice & Tense", "Prose", "Narrator voice, tense and narrative stance.", voice_tense.analyze),
        "characters": ("Characters & Relationships", "Story", "Presence, relationships and character movement.", characters.analyze),
        "continuity": ("Continuity & Timeline", "Story", "Chronology, recurring facts and inconsistencies.", continuity.analyze),
        "themes": ("Themes & Emotional Arc", "Story", "Themes and emotional movement.", themes.analyze),
        "editor": ("Editorial Patterns", "Editorial", "Clarity, redundancy and editorial signals.", editor.analyze),
        "repetition": ("Repetition & Echoes", "Prose", "Repeated words and phrases.", None),
        "pacing": ("Pacing & Momentum", "Structure", "Acceleration, slowing and sentence/paragraph movement.", None),
        "structure": ("Structure & Chapter Purpose", "Structure", "Openings, endings, paragraph shape and structural signals.", None),
        "memoir": ("Memoir Lens", "Memoir", "Reflection, event balance and memory uncertainty.", None),
        "reader": ("Reader Experience", "Editorial", "Opening, dialogue and possible reader-friction signals.", None),
        "research": ("Research & Fact Flags", "Accuracy", "Dates and claims worth checking.", None),
        "cadence": ("Cadence & Rhythm", "Prose", "Sentence movement, pauses and contrast.", cadence.analyze),
        "motifs": ("Motifs & Echoes", "Story", "Recurring words/phrases as candidate motifs.", None),
        "anchors": ("Structural Anchors", "Structure", "Recurring openings, endings and textual anchors.", None),
        "voice_dna": ("Voice DNA", "Writer", "Compare against approved personal writing samples.", voice_dna.analyze),
        "reader_perception": ("Reader Perception", "Writer", "Evidence-first impression of narrator and characters.", reader_perception.analyze),
    }


def _run_tool(key, text, all_files):
    """Dispatch to the right analyzer."""
    tools = _get_tools_dict()
    meta = tools[key]
    fn = meta[3]
    if key == "reader_perception":
        return reader_perception.analyze(text)
    if key == "voice_dna":
        return voice_dna.analyze(text)
    if fn:
        if key == "characters":
            return fn(text, all_files=all_files)
        return fn(text)
    return suite_run(key, text)


def _safe_name(n):
    n = Path(str(n or "untitled.txt")).name
    n = __import__("re").sub(r"[^A-Za-z0-9._ -]+", "_", n).strip(" .") or "untitled.txt"
    return n if Path(n).suffix.lower() in (".txt", ".md", ".text") else n + ".txt"


def _unique_path(folder, name):
    p = folder / name
    if not p.exists():
        return p
    for i in range(2, 10000):
        q = folder / f"{p.stem} ({i}){p.suffix}"
        if not q.exists():
            return q
    raise RuntimeError("Unable to create a unique filename")


def _js_safe(v):
    """Recursively convert to JSON-safe primitives."""
    if isinstance(v, dict):
        return {str(k): _js_safe(x) for k, x in v.items()}
    if isinstance(v, (list, tuple)):
        return [_js_safe(x) for x in v]
    if isinstance(v, (str, int, float, bool)) or v is None:
        return v
    if isinstance(v, Path):
        return str(v)
    return str(v)
