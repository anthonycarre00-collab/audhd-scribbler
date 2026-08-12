#!/usr/bin/env python3
"""Export module for The Audhd Scribbler. Exports never overwrite silently."""
import re
from pathlib import Path
from typing import Dict
from datetime import datetime
from .config import PROJECT_ROOT
from .file_io import read_text_file, write_text_file
from . import safety

def _output(path): return safety.unique_output_path(Path(path))

def export_markdown(file_path: str, output_path: str = None) -> str:
    path=Path(file_path)
    if not path.exists(): raise FileNotFoundError(f"File not found: {file_path}")
    if output_path is None: output_path=PROJECT_ROOT/"data"/"exports"/f"{path.stem}.md"
    out=_output(output_path); out.parent.mkdir(parents=True,exist_ok=True); write_text_file(out,read_text_file(path)); return str(out)

def export_plain_text(file_path: str, output_path: str = None) -> str:
    path=Path(file_path)
    if not path.exists(): raise FileNotFoundError(f"File not found: {file_path}")
    content=read_text_file(path)
    if content.startswith("---"):
        end=content.find("---",3)
        if end!=-1: content=content[end+3:].strip()
    content=re.sub(r'<!-- SCRIBBLER SUMMARY[\s\S]*?-->','',content).strip()
    if output_path is None: output_path=PROJECT_ROOT/"data"/"exports"/f"{path.stem}.txt"
    out=_output(output_path); out.parent.mkdir(parents=True,exist_ok=True); write_text_file(out,content); return str(out)

def _sanitize_for_docx(text):
    cleaned=[]
    for c in text:
        n=ord(c)
        if n==0: continue
        cleaned.append(c if n>=32 or n in (9,10,13) else ' ')
    return re.sub(r' {3,}','  ',''.join(cleaned))

def export_docx(file_path: str, output_path: str = None) -> str:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError: raise ImportError("python-docx is required for DOCX export")
    path=Path(file_path)
    if not path.exists(): raise FileNotFoundError(f"File not found: {file_path}")
    content=read_text_file(path); body=content
    if body.startswith("---"):
        end=body.find("---",3)
        if end!=-1: body=body[end+3:].strip()
    body=re.sub(r'<!-- SCRIBBLER SUMMARY[\s\S]*?-->','',body).strip(); body=_sanitize_for_docx(body)
    if output_path is None: output_path=PROJECT_ROOT/"data"/"exports"/f"{path.stem}.docx"
    out=_output(output_path); out.parent.mkdir(parents=True,exist_ok=True)
    doc=Document(); doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(11)
    doc.add_heading(_sanitize_for_docx(path.stem.replace('-',' ').replace('_',' ').title()),level=1)
    for para in re.split(r'\n\s*\n',body):
        para=para.strip()
        if not para: continue
        if para.startswith('# '): doc.add_heading(_sanitize_for_docx(para[2:]),level=1)
        elif para.startswith('## '): doc.add_heading(_sanitize_for_docx(para[3:]),level=2)
        elif para.startswith('### '): doc.add_heading(_sanitize_for_docx(para[4:]),level=3)
        else: doc.add_paragraph(_sanitize_for_docx(para))
    doc.save(str(out)); return str(out)

def export_analysis_report(file_path: str, analysis_results: Dict, output_path: str = None) -> str:
    path=Path(file_path)
    if output_path is None: output_path=PROJECT_ROOT/"data"/"reports"/f"{path.stem}_analysis.md"
    out=_output(output_path); out.parent.mkdir(parents=True,exist_ok=True)
    lines=[f"# Analysis Report: {path.name}",f"\n*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n"]
    for kind,result in analysis_results.items():
        lines += [f"\n---\n\n## {kind.title()}\n"]
        if isinstance(result,dict):
            if "summary" in result: lines.append(f"\n{result['summary']}\n")
            if "strengths" in result:
                lines.append("\n### Strengths\n"); lines.extend(f"- {s}" for s in result['strengths'])
            if "observations" in result:
                lines.append("\n### Observations\n")
                for obs in result['observations']:
                    if isinstance(obs,dict): lines += [f"\n**{obs.get('category','').replace('_',' ').title()}** ({obs.get('location','')})",f"\n{obs.get('formatted','')}\n"]
                    else: lines.append(f"\n- {obs}")
            for key,val in result.items():
                if key in {'summary','strengths','observations','error'}: continue
                lines.append(f"\n### {key.replace('_',' ').title()}\n")
                if isinstance(val,dict): lines.extend(f"- **{k}**: {v}" for k,v in val.items())
                elif isinstance(val,list): lines.extend(f"- {item}" for item in val)
                else: lines.append(str(val))
    write_text_file(out,'\n'.join(lines)); return str(out)
